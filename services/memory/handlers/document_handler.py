# 📄 文档处理器
import base64
import logging
import struct
import tempfile
import os
from typing import Any

from services.memory.handlers.base import BaseHandler

logger = logging.getLogger(__name__)


class DocumentHandler(BaseHandler):
    """文档模态处理器

    处理文档输入，支持：
    - PDF 解析
    - Word 文档解析
    - 纯文本提取
    - 文档摘要生成

    支持的输入格式：
    - bytes: 文档二进制数据
    - str: Base64 编码的文档数据
    """

    # 支持的文档格式
    SUPPORTED_FORMATS = {
        ".pdf": "pdf",
        ".doc": "word",
        ".docx": "word",
        ".txt": "text",
        ".md": "text",
    }

    @property
    def supported_modality(self) -> str:
        return "document"

    async def preprocess(self, content: Any, filename: str = "document.pdf") -> str:
        """预处理文档内容

        1. 解析文档提取文本
        2. 生成文档摘要

        Args:
            content: 文档数据（bytes 或 Base64 字符串）
            filename: 原始文件名（用于判断格式）

        Returns:
            str: 文档的文本描述
        """
        # 转换为 bytes
        if isinstance(content, str):
            doc_bytes = base64.b64decode(content)
        elif isinstance(content, bytes):
            doc_bytes = content
        else:
            raise ValueError(f"不支持的文档格式: {type(content)}")

        # 获取文档类型
        ext = os.path.splitext(filename)[1].lower()
        doc_type = self.SUPPORTED_FORMATS.get(ext, "text")

        # 提取文本
        if doc_type == "pdf":
            text = await self._extract_pdf_text(doc_bytes)
        elif doc_type == "word":
            text = await self._extract_word_text(doc_bytes, ext)
        else:
            text = doc_bytes.decode("utf-8", errors="ignore")

        if not text:
            return "[文档内容无法解析]"

        # 生成文档摘要
        summary = await self._generate_document_summary(text, filename)

        return summary

    async def _extract_pdf_text(self, pdf_bytes: bytes) -> str:
        """提取 PDF 文本

        Args:
            pdf_bytes: PDF 二进制数据

        Returns:
            str: 提取的文本
        """
        text = ""

        try:
            # 尝试使用 PyPDF2
            import PyPDF2
            import io

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            logger.info(f"PyPDF2 提取 PDF 文本成功: {len(text)} 字符")

        except ImportError:
            logger.warning("PyPDF2 未安装，尝试使用 pdfplumber")

            try:
                import pdfplumber
                import io

                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                logger.info(f"pdfplumber 提取 PDF 文本成功: {len(text)} 字符")

            except ImportError:
                logger.warning("pdfplumber 也未安装")

        except Exception as e:
            logger.error(f"PDF 文本提取失败: {e}")

        return text.strip()

    async def _extract_word_text(self, doc_bytes: bytes, ext: str) -> str:
        """提取 Word 文档文本

        Args:
            doc_bytes: Word 文档二进制数据
            ext: 文件扩展名

        Returns:
            str: 提取的文本
        """
        text = ""

        try:
            if ext == ".docx":
                # 使用 python-docx 解析 .docx
                from docx import Document
                import io

                doc = Document(io.BytesIO(doc_bytes))

                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"

                logger.info(f"python-docx 提取 Word 文本成功: {len(text)} 字符")

            else:
                # .doc 格式，尝试使用 antiword 或其他工具
                text = await self._extract_doc_text(doc_bytes)

        except ImportError:
            logger.warning("python-docx 未安装")

            # 尝试使用其他方法
            text = await self._extract_doc_text(doc_bytes)

        except Exception as e:
            logger.error(f"Word 文档文本提取失败: {e}")

        return text.strip()

    async def _extract_doc_text(self, doc_bytes: bytes) -> str:
        """提取 .doc 格式文本

        .doc 是旧版 Word 格式，需要特殊处理

        Args:
            doc_bytes: .doc 文档二进制数据

        Returns:
            str: 提取的文本
        """
        try:
            import asyncio

            # 写入临时文件
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(doc_bytes)
                tmp_path = tmp.name

            try:
                # 尝试使用 antiword（Linux/Mac）
                process = await asyncio.create_subprocess_exec(
                    "antiword", tmp_path,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                )
                stdout, _ = await process.communicate()

                if process.returncode == 0:
                    return stdout.decode("utf-8", errors="ignore")

            except FileNotFoundError:
                pass

            finally:
                os.unlink(tmp_path)

        except Exception as e:
            logger.error(f".doc 文本提取失败: {e}")

        # 如果 antiword 不可用，尝试提取文本
        return doc_bytes.decode("utf-8", errors="ignore")

    async def _generate_document_summary(self, text: str, filename: str) -> str:
        """生成文档摘要

        Args:
            text: 文档文本
            filename: 文件名

        Returns:
            str: 文档摘要
        """
        # 截取前 4000 字符用于摘要
        sample_text = text[:4000]
        if len(text) > 4000:
            sample_text += "\n...[内容已截断]..."

        prompt = f"""请为以下文档生成摘要。

文件名: {filename}
文档内容:
{sample_text}

请提供：
1. 文档主题
2. 主要内容概述
3. 关键要点（如有）"""

        try:
            summary = await self.llm.generate_chat_response(
                system_prompt="你是一个文档分析专家，擅长总结文档内容。",
                context="",
                user_query=prompt,
            )
            return summary

        except Exception as e:
            logger.error(f"文档摘要生成失败: {e}")
            # 返回原文片段
            return f"文档内容：\n{text[:500]}..."

    async def get_vector(self, text: str) -> bytes:
        """生成向量"""
        embedding = await self.llm.get_embedding(text)
        return struct.pack(f'{len(embedding)}f', *embedding)

    async def store_raw_content(self, content: Any, filename: str = "document.pdf") -> str:
        """存储文档到 OSS

        Args:
            content: 文档数据（bytes）
            filename: 原始文件名

        Returns:
            str: OSS 存储路径
        """
        if not self.oss_client:
            raise ValueError("OSS 客户端未配置")

        if not self.user_id:
            raise ValueError("user_id 未设置")

        if isinstance(content, str):
            doc_bytes = base64.b64decode(content)
        elif isinstance(content, bytes):
            doc_bytes = content
        else:
            raise ValueError(f"不支持的文档格式: {type(content)}")

        # 上传到 OSS
        path = await self.oss_client.upload(doc_bytes, filename, "document", self.user_id)
        logger.info(f"文档存储成功: {path}")

        return path
